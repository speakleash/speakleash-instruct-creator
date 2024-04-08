from datasets import load_dataset
import docx
import json
import os
from utils.functions import get_dir_path

# This is not a complete script to translate and generate instructions.
# This part only creates small .docx files from the dataset, which can then be conveniently translated using external services.
# Later create a dataset from the translated documents using orca_math_create_json_from_docx.py

dataset = load_dataset("microsoft/orca-math-word-problems-200k")
questions, answers = dataset['train']['question'], dataset['train']['answer']
base_dir = os.path.dirname(os.path.abspath(__file__))
data_folder = get_dir_path("data") or os.path.join(base_dir, "data")
translated_folder = os.path.join(data_folder, "translated")
indexes_path = os.path.join(data_folder, 'docx_indexes.json')


def split_dataset_into_indexes(question_arr, answer_arr, max_chars: int = 4800000) -> list:
    """
    Args:
        question_arr: array with all questions
        answer_arr: array with all answers
        max_chars: maximum number of characters in single split

    Returns: array indices between which the sum of all questions and answers (including margins)
    is less than max_chars.
    """
    subset_indexes = []
    current_chars_count = 0
    start_index = 0
    entry_margin = 12

    for index, (question, answer) in enumerate(zip(question_arr, answer_arr)):
        entry_length = len(question) + len(answer) + entry_margin

        if current_chars_count + entry_length > max_chars:
            subset_indexes.append((start_index, index - 1))
            start_index = index
            current_chars_count = 0

        current_chars_count += entry_length

    if current_chars_count > 0:
        subset_indexes.append((start_index, len(question_arr)))

    return subset_indexes


indexes = split_dataset_into_indexes(questions, answers)


def save_indexes(file_path: str) -> None:
    with open(file_path, 'w') as file:
        json.dump(indexes, file)


os.makedirs(data_folder, exist_ok=True)
os.makedirs(translated_folder, exist_ok=True)
save_indexes(indexes_path)

print("Indexes are generated and stored in the folder " + indexes_path)

def generate_docx(questions, answers):
    doc = docx.Document()
    for row_idx in range(start, end):
        doc.add_paragraph("#Q#\n")
        doc.add_paragraph(questions[row_idx] + "\n")
        doc.add_paragraph("#A#\n")
        doc.add_paragraph(answers[row_idx] + "\n")
    doc.save(os.path.join(data_folder, f"math_problems_to_translate_char_limit{doc_idx}.docx"))


for doc_idx, (start, end) in enumerate(indexes):
    generate_docx(questions, answers)

print(".docx with questions and answers are created in " + data_folder)
print("Now translate them and put into " + translated_folder)
