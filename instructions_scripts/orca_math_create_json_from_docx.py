import re
import json
import os

from datasets import load_dataset
from docx import Document

from utils.functions import get_dir_path

# This is not a complete script to translate and generate instructions.
# You need to generate .docx documents using orca_math_create_english_docx.py,
# translate them in an external service, place them in /data/translated
# and finally use this script to create a dataset from them.

dataset = load_dataset("microsoft/orca-math-word-problems-200k")
questions, answers = dataset['train']['question'], dataset['train']['answer']
base_dir = os.path.dirname(os.path.abspath(__file__))
data_folder = get_dir_path("data") or os.path.join(base_dir, "data")
translated_folder = os.path.join(data_folder, "translated")
indexes_path = os.path.join(data_folder, 'docx_indexes.json')
translated_json_path = os.path.join(translated_folder, 'json')


def load_indexes(file_path):
    with open(file_path, 'r') as file:
        return json.load(file)


def remove_trailing_newline(text):
    if text.endswith('\n'):
        return text[:-1]
    return text


indexes = load_indexes(indexes_path)
print("Indexes are loaded.")

paths = os.listdir(translated_folder)
paths = [path for path in paths if path.endswith(".docx")]
paths = sorted(paths, key=lambda x: int(re.search(r'\d+', x).group()))


os.makedirs(translated_json_path, exist_ok=True)

for doc_idx in range(0, len(indexes)):
    print("Create .json file from document number: " + str(doc_idx))
    # Load the DOCX document
    doc_path = os.path.join(translated_folder, paths[doc_idx])
    doc = Document(doc_path)

    # Initialize variables
    data = []
    id_counter = 1

    start_idx = indexes[doc_idx][0]
    # Iterate through the paragraphs in the document
    for i in range(0, len(doc.paragraphs), 4):
        if doc.paragraphs[i].text == "#Q#\n" and doc.paragraphs[i + 2].text == "#A#\n":
            # Create a dictionary for the question-answer set
            entry = {
                "instruct": remove_trailing_newline(doc.paragraphs[i + 1].text),
                "input": "",
                "output": remove_trailing_newline(doc.paragraphs[i + 3].text),
                "instruct_eng": questions[start_idx + i // 4],
                "output_eng": answers[start_idx + i // 4],
                "source_name": "microsoft-orca_math_problems",
                "source_url": "https://huggingface.co/datasets/microsoft/orca-math-word-problems-200k",
                "source_description": "Pary pytanie-odpowiedź powstały na bazie zestawu danych orca math word problems, a następnie przetłumaczone maszynowo.",
                "script_name": "translate_orca_math_problems.py",
                "id": start_idx + i // 4,
                "status": "",
                "updated_by": ""
            }
            # Append the entry to the data list
            data.append(entry)
        else:
            print("something went wrong for entry " + str(i // 4))



    # Save the data to a JSONL file
    jsonl_path = os.path.join(translated_json_path, f"math_problems_converted{doc_idx}.jsonl")
    with open(jsonl_path, 'w', encoding='utf-8-sig') as outfile:
        for entry in data:
            json.dump(entry, outfile, ensure_ascii=False)
            outfile.write('\n')

print("All .docx documents have been converted to .json")

paths = os.listdir(translated_json_path)
paths = sorted(paths, key=lambda x: int(re.search(r'\d+', x).group()))

base_dir = os.path.dirname(os.path.abspath(__file__))
output_path = get_dir_path("output") or os.path.join(base_dir, "output")
combined_file_path = os.path.join(output_path, 'polish_orca_math.json')
with open(combined_file_path, 'w') as output_file:
    output_file.write("[\n")

    first_object = True

    for path in paths:
        with open(os.path.join(translated_json_path, path), 'r', encoding='utf-8-sig') as input_file:
            for line in input_file:
                if line.strip():
                    if first_object:
                        first_object = False
                    else:
                        output_file.write(',')
                        output_file.write("\n")

                    json_object = json.loads(line)
                    json.dump(json_object, output_file, ensure_ascii=False, indent=4)

    output_file.write('\n]')

print("Combined file saved into " + combined_file_path)
print("You can now delete all other files created during translation.")
